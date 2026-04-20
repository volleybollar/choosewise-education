#!/usr/bin/env python3
"""Build the NotebookLM 140 Styles — English Prompt Document.

Produces three artifacts from the translated prompt JSON:
  1. exports/nlm-140-prompts-en.html (editable HTML source, A4-print-ready)
  2. assets/pdfs/nlm-140-prompts-en.pdf (PDF for sharing/print)
  3. exports/nlm-140-prompts-en.docx (editable Word document)

Input: /tmp/nlm-prompts-en-chunk{1..4}.json — translated prompt chunks
       (each covers 35 prompts: 1-35, 36-70, 71-105, 106-140).

Run: python3 exports/build-nlm-prompts-en.py
"""
import json
import html
from pathlib import Path
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(__file__).parent.parent
html_out = root / "exports/nlm-140-prompts-en.html"
pdf_out = root / "assets/pdfs/nlm-140-prompts-en.pdf"
docx_out = root / "exports/nlm-140-prompts-en.docx"

# ─── 1. Merge the 4 translated chunks ──────────────────────────────
prompts: dict[int, str] = {}
for idx in range(1, 5):
    src = Path(f"/tmp/nlm-prompts-en-chunk{idx}.json")
    if not src.exists():
        raise SystemExit(f"Missing chunk: {src}")
    data = json.loads(src.read_text())
    for k, v in data.items():
        prompts[int(k)] = v

missing = [n for n in range(1, 141) if n not in prompts]
if missing:
    raise SystemExit(f"Missing prompt numbers: {missing}")
print(f"Merged {len(prompts)} prompts from 4 chunks.")

# ─── 2. Render HTML source ─────────────────────────────────────────
HTML_HEAD = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>NotebookLM 140 Styles — Prompt Library (English)</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700&family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
<style>
  @page { size: A4; margin: 18mm 18mm 20mm 18mm; }
  * { margin: 0; padding: 0; box-sizing: border-box; }

  body {
    font-family: 'Inter', -apple-system, sans-serif;
    color: #1B2733;
    background: #ffffff;
    -webkit-print-color-adjust: exact;
    print-color-adjust: exact;
    font-size: 10.5pt;
    line-height: 1.55;
  }

  /* ─── Cover page ─── */
  .cover {
    page-break-after: always;
    min-height: 260mm;
    display: flex;
    flex-direction: column;
    justify-content: center;
    background: #F6F3EE;
    padding: 40mm 30mm;
    margin: -18mm -18mm 0 -18mm;   /* break out of @page margin */
  }
  .cover .eyebrow {
    font-size: 10pt;
    font-weight: 700;
    letter-spacing: 3px;
    text-transform: uppercase;
    color: #C8A86B;
    margin-bottom: 18px;
  }
  .cover h1 {
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 44pt;
    font-weight: 700;
    line-height: 1.08;
    letter-spacing: -1.2px;
    color: #1B2733;
    margin-bottom: 14px;
    max-width: 140mm;
  }
  .cover .subtitle {
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 18pt;
    font-style: italic;
    color: #3E4A57;
    margin-bottom: 28px;
    max-width: 150mm;
  }
  .cover .lede {
    font-size: 11.5pt;
    line-height: 1.65;
    color: #3E4A57;
    max-width: 140mm;
    margin-bottom: 42px;
  }
  .cover .accent-rule {
    width: 80px;
    height: 3px;
    background: #C8A86B;
    margin-bottom: 18px;
  }
  .cover .byline {
    font-size: 10pt;
    color: #3E4A57;
    margin-top: auto;
  }
  .cover .byline strong {
    font-weight: 700;
    color: #1B2733;
  }

  /* ─── Intro page ─── */
  .intro {
    page-break-after: always;
    padding-top: 4mm;
  }
  .intro h2 {
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 22pt;
    font-weight: 700;
    color: #1B2733;
    margin-bottom: 10px;
    letter-spacing: -0.4px;
  }
  .intro .eyebrow {
    font-size: 9pt;
    font-weight: 700;
    letter-spacing: 2.4px;
    text-transform: uppercase;
    color: #C8A86B;
    margin-bottom: 6px;
  }
  .intro p {
    font-size: 10.5pt;
    line-height: 1.7;
    color: #3E4A57;
    margin-bottom: 10px;
    max-width: 160mm;
  }
  .intro .notes {
    margin-top: 22px;
    padding: 14px 18px;
    background: #F6F3EE;
    border-left: 3px solid #C8A86B;
    border-radius: 2px;
  }
  .intro .notes p { margin-bottom: 6px; }
  .intro .notes p:last-child { margin-bottom: 0; }

  /* ─── Prompt block ─── */
  .prompt {
    page-break-inside: avoid;
    margin-bottom: 14pt;
    padding-bottom: 10pt;
    border-bottom: 1px solid #E8E4DC;
  }
  .prompt:last-child {
    border-bottom: none;
  }
  .prompt .meta {
    display: flex;
    align-items: baseline;
    gap: 10px;
    margin-bottom: 4pt;
  }
  .prompt .num {
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 22pt;
    font-weight: 700;
    line-height: 1;
    color: #C8A86B;
    min-width: 42px;
  }
  .prompt .label {
    font-size: 8.5pt;
    font-weight: 700;
    letter-spacing: 1.6px;
    text-transform: uppercase;
    color: #3E4A57;
  }
  .prompt .body {
    font-size: 10pt;
    line-height: 1.6;
    color: #1B2733;
    text-align: justify;
    hyphens: auto;
  }

  /* Header/footer styling — Playwright will inject footer template */
  .page-break { page-break-after: always; }
</style>
</head>
<body>
"""

HTML_COVER = """
<section class="cover">
  <div class="accent-rule"></div>
  <p class="eyebrow">NotebookLM Style Library</p>
  <h1>140 Prompts for AI Infographics</h1>
  <p class="subtitle">English edition — every style, every directive, fully translated from the Swedish source.</p>
  <p class="lede">This document collects all 140 art-direction prompts that produce the NotebookLM infographic styles on choosewise.education. Each prompt is numbered to match the original Swedish library and can be pasted directly into NotebookLM, Midjourney, Flux, or any other image generator that accepts long-form directives.</p>
  <p class="byline"><strong>Johan Lindström</strong> · choosewise.education · Updated April 2026</p>
</section>
"""

HTML_INTRO = """
<section class="intro">
  <p class="eyebrow">About this library</p>
  <h2>How to use these prompts</h2>
  <p>Each numbered entry below is a complete, stand-alone prompt. Copy the prompt in full and paste it into the image-generation field of your chosen tool. The numbering matches the Swedish library, so the same number produces the same style across languages — useful when you want to A/B the two wordings.</p>
  <p>The prompts intentionally preserve English terms that were already anglicised in the Swedish source (glassmorphism, cel-shading, bokeh, HUD, bento grid, etc.) and all colour codes, font sizes, and proportions are kept as-is.</p>
  <div class="notes">
    <p><strong>Licence.</strong> These prompts are © Johan Lindström, shared under CC BY-NC-SA 4.0. You may use, share, and adapt them for non-commercial purposes with attribution.</p>
    <p><strong>Tip.</strong> If a generator truncates long prompts, split the prompt at a sentence boundary and send it in two paragraphs — the style tends to survive the split.</p>
  </div>
</section>
"""

HTML_TAIL = "</body></html>\n"

blocks = [HTML_HEAD, HTML_COVER, HTML_INTRO]
for n in range(1, 141):
    body = html.escape(prompts[n])
    blocks.append(f"""
<div class="prompt">
  <div class="meta">
    <span class="num">{n}</span>
    <span class="label">Style {n:03d}</span>
  </div>
  <p class="body">{body}</p>
</div>""")
blocks.append(HTML_TAIL)

html_doc = "\n".join(blocks)
html_out.write_text(html_doc, encoding="utf-8")
print(f"HTML source: {html_out} ({html_out.stat().st_size // 1024} KB)")

# ─── 3. Build PDF via Playwright ──────────────────────────────────
FOOTER_HTML = """
<div style="font-size: 8pt; color: #8a8a85; width: 100%; padding: 0 18mm;
            display: flex; justify-content: space-between; align-items: center;
            font-family: 'Inter', -apple-system, sans-serif;">
  <span>NotebookLM 140 Styles · Johan Lindström · choosewise.education</span>
  <span>Page <span class="pageNumber"></span> / <span class="totalPages"></span></span>
</div>
"""

EMPTY_HEADER = "<div></div>"

with sync_playwright() as p:
    browser = p.chromium.launch()
    page = browser.new_page()
    page.goto(html_out.as_uri(), wait_until="networkidle")
    page.emulate_media(media="print")
    pdf_out.parent.mkdir(parents=True, exist_ok=True)
    page.pdf(
        path=str(pdf_out),
        format="A4",
        print_background=True,
        margin={"top": "0", "bottom": "14mm", "left": "0", "right": "0"},
        prefer_css_page_size=True,
        display_header_footer=True,
        header_template=EMPTY_HEADER,
        footer_template=FOOTER_HTML,
    )
    browser.close()
print(f"PDF: {pdf_out} ({pdf_out.stat().st_size // 1024} KB)")

# ─── 4. Build DOCX via python-docx ────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)

# Default body font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Cover
title = doc.add_paragraph()
run = title.add_run("NotebookLM Style Library")
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0xC8, 0xA8, 0x6B)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

h1 = doc.add_paragraph()
r = h1.add_run("140 Prompts for AI Infographics")
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1B, 0x27, 0x33)

subtitle = doc.add_paragraph()
r = subtitle.add_run("English edition — every style, every directive, fully translated from the Swedish source.")
r.font.size = Pt(14)
r.font.italic = True
r.font.color.rgb = RGBColor(0x3E, 0x4A, 0x57)

doc.add_paragraph()
intro = doc.add_paragraph(
    "This document collects all 140 art-direction prompts that produce the "
    "NotebookLM infographic styles on choosewise.education. Each prompt is "
    "numbered to match the original Swedish library and can be pasted directly "
    "into NotebookLM, Midjourney, Flux, or any other image generator that "
    "accepts long-form directives."
)

doc.add_paragraph()
byline = doc.add_paragraph()
r = byline.add_run("Johan Lindström · choosewise.education · Updated April 2026")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x3E, 0x4A, 0x57)

doc.add_page_break()

# How-to section
h2 = doc.add_paragraph()
r = h2.add_run("How to use these prompts")
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1B, 0x27, 0x33)

doc.add_paragraph(
    "Each numbered entry below is a complete, stand-alone prompt. Copy the "
    "prompt in full and paste it into the image-generation field of your "
    "chosen tool. The numbering matches the Swedish library, so the same "
    "number produces the same style across languages — useful when you want "
    "to A/B the two wordings."
)
doc.add_paragraph(
    "The prompts intentionally preserve English terms that were already "
    "anglicised in the Swedish source (glassmorphism, cel-shading, bokeh, "
    "HUD, bento grid, etc.) and all colour codes, font sizes, and "
    "proportions are kept as-is."
)

p = doc.add_paragraph()
r = p.add_run("Licence. ")
r.font.bold = True
p.add_run(
    "These prompts are © Johan Lindström, shared under CC BY-NC-SA 4.0. "
    "You may use, share, and adapt them for non-commercial purposes with "
    "attribution."
)

p = doc.add_paragraph()
r = p.add_run("Tip. ")
r.font.bold = True
p.add_run(
    "If a generator truncates long prompts, split the prompt at a sentence "
    "boundary and send it in two paragraphs — the style tends to survive "
    "the split."
)

doc.add_page_break()

# The 140 prompts
for n in range(1, 141):
    heading = doc.add_paragraph()
    r = heading.add_run(f"Style {n:03d}")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xC8, 0xA8, 0x6B)

    body_para = doc.add_paragraph(prompts[n])
    body_para.paragraph_format.space_after = Pt(14)

doc.save(docx_out)
print(f"DOCX: {docx_out} ({docx_out.stat().st_size // 1024} KB)")

print("Done.")
